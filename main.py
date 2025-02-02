import os
from docx import Document

def list_files(directory, extension):
    """
    Récupère tous les fichiers avec une certaine extension dans un répertoire.
    """
    file_list = []
    for root, dirs, files in os.walk(directory):
        for file in files:
            if file.endswith(extension):
                file_list.append(os.path.join(root, file))
    return file_list

def create_word_file(pdf_files, output_file):
    """
    Crée un fichier Word contenant les noms des fichiers PDF.
    """
    document = Document()
    document.add_heading('Liste des fichiers PDF', level=1)

    for pdf_file in pdf_files:
        document.add_paragraph(os.path.basename(pdf_file))

    document.save(output_file)
    print(f"Le fichier Word '{output_file}' a été créé avec succès.")

# Spécifiez le chemin du dossier contenant les fichiers PDF
dossier_pdf = r"D:\livre"

# Spécifiez le chemin de sortie pour le fichier Word
fichier_word_sortie = "liste_pdf.docx"

# Récupère la liste des fichiers PDF dans le dossier spécifié
pdf_files = list_files(dossier_pdf, ".pdf")

# Crée le fichier Word avec la liste des noms de fichiers PDF
create_word_file(pdf_files, fichier_word_sortie)
